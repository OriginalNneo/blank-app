import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import json
import os
from datetime import datetime
from typing import List, Dict, Any, Optional
from app.utils.config import get_gemini_api_key
from app.services.attendance_service import AttendanceService
from pydantic import BaseModel


class MeetingMinutesRequest(BaseModel):
    meeting_title: str = "Corporate Board Meeting"
    date_time: Optional[str] = None
    company: Optional[str] = None
    location: Optional[str] = None
    attendees: Optional[str] = None
    absent: Optional[str] = None
    meeting_chair: Optional[str] = None


class MinutesService:
    def __init__(self):
        # Get API key from configuration
        api_key = get_gemini_api_key()

        if not api_key:
            raise ValueError("Gemini API key not found in configuration")

        genai.configure(api_key=api_key)
        self._model = None  # Lazy initialization
        self._api_key = api_key

    def _initialize_gemini_model(self):
        """Initialize Gemini model with preferred settings"""
        try:
            # Try preferred models in order
            preferred_models = ['gemini-2.0-flash-exp', 'gemini-2.0-flash', 'gemini-1.5-pro', 'gemini-1.5-flash']

            for model_name in preferred_models:
                try:
                    model = genai.GenerativeModel(model_name)
                    print(f"Using Gemini model: {model_name}")
                    return model
                except Exception as e:
                    print(f"Model {model_name} failed: {e}")
                    continue

            # Fallback to any available model
            models = genai.list_models()
            for model_info in models:
                if 'generateContent' in model_info.supported_generation_methods:
                    try:
                        model = genai.GenerativeModel(model_info.name)
                        print(f"Using fallback model: {model_info.name}")
                        return model
                    except:
                        continue

            raise Exception("No suitable Gemini model found")

        except Exception as e:
            raise Exception(f"Failed to initialize Gemini model: {e}")


    def _get_model(self):
        """Get or initialize the Gemini model (lazy initialization)"""
        if self._model is None:
            print("Initializing Gemini model...")
            self._model = self._initialize_gemini_model()
            print("Gemini model initialized successfully")
        return self._model

    def process_content_with_gemini(self, content: str) -> Dict[str, Any]:
        """Process meeting content with Gemini to extract structured information"""
        try:
            # Get model (will initialize if needed)
            model = self._get_model()
            
            prompt = """
            You are an expert at analyzing meeting information and extracting structured data for meeting minutes.

            Analyze the following meeting content and extract:
            1. Meeting title/type
            2. Agenda items (list of topics discussed)
            3. Key discussion points for each agenda item
            4. Decisions made
            5. Action items (if any)
            6. Any dates, locations, or other relevant details mentioned

            Return ONLY valid JSON with this exact structure:
            {
                "meeting_title": "Meeting title or type",
                "agenda_items": [
                    {
                        "item_number": 1,
                        "title": "Agenda item title (e.g., 'Call to Order', 'Approval of Agenda', 'Chair's Report')",
                        "description": "Detailed description of what was discussed or decided for this agenda item",
                        "action_items": ["Action item 1", "Action item 2"]
                    }
                ],
                "extracted_date": "Date mentioned in presentation (if any)",
                "extracted_location": "Location mentioned (if any)",
                "extracted_company": "Company/organization name (if any)"
            }

            Rules:
            1. Standardize agenda items to common meeting formats (Call to Order, Approval of Agenda, Approval of Previous Meeting Minutes, Chair's Report, CEO's Report, Committee Reports, Old Business, New Business, Other Business, Adjournment)
            2. If agenda items are not explicitly stated, infer them from the content structure
            3. Provide detailed descriptions for each agenda item based on the presentation content
            4. Extract action items if mentioned
            5. Return ONLY the JSON, no other text
            """

            # Process the content
            print(f"Processing {len(content)} characters with Gemini...")
            response = model.generate_content([
                prompt,
                content
            ])
            print("Got response from Gemini")

            # Parse response
            response_text = response.text.strip()

            # Clean the response text
            if response_text.startswith("```json"):
                response_text = response_text[7:]
            if response_text.startswith("```"):
                response_text = response_text[3:]
            if response_text.endswith("```"):
                response_text = response_text[:-3]

            response_text = response_text.strip()

            # Find JSON in response
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1

            if json_start != -1 and json_end > json_start:
                response_text = response_text[json_start:json_end]

            # Parse JSON
            try:
                parsed_data = json.loads(response_text)
                return parsed_data
            except json.JSONDecodeError as json_err:
                print(f"JSON decode error: {json_err}")
                print(f"Response text that failed to parse: {response_text[:500]}")
                # Return a default structure if JSON parsing fails
                return {
                    "meeting_title": "Meeting",
                    "agenda_items": [
                        {
                            "item_number": 1,
                            "title": "General Discussion",
                            "description": powerpoint_text[:500] if len(powerpoint_text) > 0 else "No content extracted",
                            "action_items": []
                        }
                    ],
                    "extracted_date": None,
                    "extracted_location": None,
                    "extracted_company": None
                }

        except Exception as e:
            print(f"Error processing content with Gemini: {e}")
            import traceback
            print(traceback.format_exc())
            raise Exception(f"Failed to process meeting content: {str(e)}")

    def generate_minutes_word(
        self, 
        request: MeetingMinutesRequest,
        processed_data: Dict[str, Any]
    ) -> bytes:
        """Generate meeting minutes Word document following the new format"""
        try:
            doc = Document()
            
            # Set page margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Add logo at the top left
            try:
                # Try multiple possible paths to find the logo
                current_file = os.path.abspath(__file__)
                possible_paths = [
                    # Path 1: Absolute path (most reliable)
                    '/Users/nathanielneo/Desktop/TGYN_Admin/Image/TGYN Logo S.jpeg',
                    # Path 2: From backend/app/services/minutes_service.py -> go up 4 levels
                    os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(current_file)))), 'Image', 'TGYN Logo S.jpeg'),
                    # Path 3: Relative from backend directory
                    os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(current_file))), '..', 'Image', 'TGYN Logo S.jpeg'),
                    # Path 4: From project root (if running from root)
                    os.path.join(os.getcwd(), 'Image', 'TGYN Logo S.jpeg'),
                    # Path 5: Try from current working directory with various parent levels
                    os.path.join(os.getcwd(), '..', 'Image', 'TGYN Logo S.jpeg'),
                ]
                
                logo_path = None
                for path in possible_paths:
                    normalized_path = os.path.normpath(path)
                    if os.path.exists(normalized_path):
                        logo_path = normalized_path
                        break
                
                if logo_path and os.path.exists(logo_path):
                    # Create a paragraph for the logo at the very beginning
                    logo_paragraph = doc.add_paragraph()
                    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Add the image to the paragraph
                    run = logo_paragraph.add_run()
                    # Add image with width of 1.5 inches (adjust as needed)
                    # Height will be calculated automatically to maintain aspect ratio
                    run.add_picture(logo_path, width=Inches(1.5))
                    print(f"Logo added successfully from: {logo_path}")
                    
                    # Add spacing after logo
                    doc.add_paragraph()
                else:
                    print(f"Logo file not found. Tried paths: {possible_paths}")
                    print(f"Current working directory: {os.getcwd()}")
                    print(f"Current file: {current_file}")
            except Exception as e:
                import traceback
                print(f"Error adding logo: {e}")
                print(f"Traceback: {traceback.format_exc()}")
                # Continue without logo if there's an error

            # Helper function to add section header with pink background
            def add_section_header(text, number=None):
                p = doc.add_paragraph()
                if number:
                    run = p.add_run(f"{number}. {text}")
                else:
                    run = p.add_run(text)
                run.font.size = Pt(12)
                run.font.name = 'Calibri'
                run.bold = True
                run.font.color.rgb = RGBColor(50, 50, 50)  # Dark gray text
                
                # Add light pink background
                pPr = p._element.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F8D7DA')  # Light pink
                pPr.append(shd)
                
                # Add spacing before and after
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '120')
                spacing.set(qn('w:after'), '120')
                pPr.append(spacing)
                
                return p

            # Parse date and time
            date_str = ""
            time_str = ""
            if request.date_time:
                try:
                    if 'T' in request.date_time:
                        date_part, time_part = request.date_time.split('T')
                        try:
                            from datetime import datetime
                            dt = datetime.strptime(date_part, "%Y-%m-%d")
                            date_str = dt.strftime("%B %d, %Y")  # e.g., "September 21, 2025"
                        except:
                            date_str = date_part
                        time_str = time_part[:5] if len(time_part) >= 5 else time_part
                    else:
                        try:
                            from datetime import datetime
                            dt = datetime.strptime(request.date_time, "%Y-%m-%d")
                            date_str = dt.strftime("%B %d, %Y")
                        except:
                            date_str = request.date_time
                except:
                    date_str = request.date_time

            # Title - Centered, large, bold
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_p.add_run(request.meeting_title or "Meeting Minutes")
            title_run.font.size = Pt(18)
            title_run.font.name = 'Calibri'
            title_run.bold = True
            title_run.font.color.rgb = RGBColor(50, 50, 50)  # Dark gray
            
            # Decorative line below title
            line_p = doc.add_paragraph()
            line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            line_run = line_p.add_run("_" * 60)
            line_run.font.size = Pt(10)
            line_run.font.color.rgb = RGBColor(100, 100, 100)
            
            doc.add_paragraph()  # Spacing

            # Attendance Section at the top
            attendance_heading = doc.add_paragraph()
            attendance_heading_run = attendance_heading.add_run("Attendance")
            attendance_heading_run.font.size = Pt(14)
            attendance_heading_run.font.name = 'Calibri'
            attendance_heading_run.bold = True
            attendance_heading_run.font.color.rgb = RGBColor(50, 50, 50)
            attendance_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add light pink background to attendance header
            pPr = attendance_heading._element.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'F8D7DA')
            pPr.append(shd)
            
            # Fetch attendance data from Google Sheets - use most recent date
            present_members = []
            absent_members = []
            
            try:
                # Get the most recent attendance from Google Sheets
                attendance_service = AttendanceService()
                members = attendance_service.get_members()
                attendance_data, attendance_date_used = attendance_service.get_most_recent_attendance()
                
                print(f"Using attendance data from date: {attendance_date_used}")
                
                # Create a mapping of name to member info (case-insensitive)
                members_map = {}
                for member in members:
                    members_map[member['name'].lower()] = member
                
                # Separate present and absent members
                for name, status in attendance_data.items():
                    name_lower = name.lower()
                    if name_lower in members_map:
                        member_info = members_map[name_lower]
                        if status.lower() == 'present':
                            present_members.append({
                                'name': member_info['name'],
                                'address': member_info.get('address', member_info['name'])
                            })
                        else:
                            absent_members.append({
                                'name': member_info['name'],
                                'address': member_info.get('address', member_info['name'])
                            })
            except Exception as e:
                print(f"Error fetching attendance from Google Sheets: {e}")
                import traceback
                print(traceback.format_exc())
                # Fallback to using request.attendees and request.absent if available
                if request.attendees:
                    for attendee in [a.strip() for a in request.attendees.split(',') if a.strip()]:
                        present_members.append({'name': attendee, 'address': attendee})
                if request.absent:
                    for absent in [a.strip() for a in request.absent.split(',') if a.strip()]:
                        absent_members.append({'name': absent, 'address': absent})
            
            # Present Section
            if present_members:
                present_heading = doc.add_paragraph()
                present_heading_run = present_heading.add_run("Present")
                present_heading_run.font.size = Pt(12)
                present_heading_run.font.name = 'Calibri'
                present_heading_run.bold = True
                present_heading_run.font.color.rgb = RGBColor(50, 50, 50)
                present_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for member in present_members:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(f"{member['address']} {member['name']}")
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
            
            # Absent with Apologies Section
            if absent_members:
                absent_heading = doc.add_paragraph()
                absent_heading_run = absent_heading.add_run("Absent with Apologies")
                absent_heading_run.font.size = Pt(12)
                absent_heading_run.font.name = 'Calibri'
                absent_heading_run.bold = True
                absent_heading_run.font.color.rgb = RGBColor(50, 50, 50)
                absent_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for member in absent_members:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(f"{member['address']} {member['name']}")
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
            
            doc.add_paragraph()  # Spacing

            # Meeting Metadata Section
            metadata_items = [
                ("Date:", date_str or ""),
                ("Time:", time_str or ""),
                ("Location:", request.location or ""),
            ]
            
            for label, value in metadata_items:
                p = doc.add_paragraph()
                run1 = p.add_run(f"{label} ")
                run1.font.size = Pt(11)
                run1.font.name = 'Calibri'
                run1.bold = True
                run2 = p.add_run(value)
                run2.font.size = Pt(11)
                run2.font.name = 'Calibri'
            
            doc.add_paragraph()  # Spacing

            # Get agenda items
            agenda_items = processed_data.get("agenda_items", [])
            if not agenda_items:
                agenda_items = [{"item_number": 1, "title": "General Discussion", "description": "Meeting discussion", "action_items": []}]

            # Section 1: Call to Order
            add_section_header("Call to Order", 1)
            chair = request.meeting_chair or "[Chairperson's Name]"
            time_display = time_str or "[Time]"
            call_to_order_p = doc.add_paragraph()
            call_to_order_run = call_to_order_p.add_run(f"The meeting was called to order by {chair} at {time_display}.")
            call_to_order_run.font.size = Pt(11)
            call_to_order_run.font.name = 'Calibri'
            
            doc.add_paragraph()  # Spacing

            # Section 2: Executive Reports / Agenda Items
            add_section_header("Executive Reports", 2)
            
            # Group agenda items or display them
            for idx, item in enumerate(agenda_items[:5]):  # Limit to first 5 items for reports
                title = item.get('title', 'General Discussion')
                description = item.get('description', '')
                action_items = item.get('action_items', [])
                
                # Sub-section title (bold)
                sub_p = doc.add_paragraph()
                sub_run = sub_p.add_run(title)
                sub_run.font.size = Pt(11)
                sub_run.font.name = 'Calibri'
                sub_run.bold = True
                
                # Description with bullet
                desc_p = doc.add_paragraph()
                desc_run = desc_p.add_run(f"• {description}")
                desc_run.font.size = Pt(11)
                desc_run.font.name = 'Calibri'
                
                # Action items if any
                if action_items:
                    for ai in action_items[:2]:  # Limit action items
                        ai_p = doc.add_paragraph()
                        ai_run = ai_p.add_run(f"• {str(ai)}")
                        ai_run.font.size = Pt(11)
                        ai_run.font.name = 'Calibri'
            
            doc.add_paragraph()  # Spacing

            # Section 3: Ongoing Issues
            add_section_header("Ongoing Issues", 3)
            ongoing_p = doc.add_paragraph()
            ongoing_run = ongoing_p.add_run("Items that were set aside in previous meetings.")
            ongoing_run.font.size = Pt(11)
            ongoing_run.font.name = 'Calibri'
            
            doc.add_paragraph()  # Spacing

            # Section 4: New Issues
            add_section_header("New Issues", 4)
            
            # Use remaining agenda items or create placeholder
            new_items = agenda_items[5:] if len(agenda_items) > 5 else []
            if not new_items:
                new_items = [{"title": "New Item 1", "description": "Discussion led by [Name].", "action_items": ["Action items and responsible persons."]},
                           {"title": "New Item 2", "description": "Discussion led by [Name].", "action_items": ["Action items and responsible persons."]}]
            
            for item in new_items[:2]:  # Limit to 2 new items
                title = item.get('title', 'New Item')
                description = item.get('description', '')
                action_items = item.get('action_items', [])
                
                # Item title (bold)
                item_p = doc.add_paragraph()
                item_run = item_p.add_run(title)
                item_run.font.size = Pt(11)
                item_run.font.name = 'Calibri'
                item_run.bold = True
                
                # Description with bullet
                desc_p = doc.add_paragraph()
                desc_run = desc_p.add_run(f"• {description}")
                desc_run.font.size = Pt(11)
                desc_run.font.name = 'Calibri'
                
                # Action items
                for ai in action_items[:2]:
                    ai_p = doc.add_paragraph()
                    ai_run = ai_p.add_run(f"• {str(ai)}")
                    ai_run.font.size = Pt(11)
                    ai_run.font.name = 'Calibri'
            
            doc.add_paragraph()  # Spacing

            # Section 5: Next Meeting
            add_section_header("Next Meeting", 5)
            next_meeting_items = [
                ("Date:", ""),
                ("Time:", ""),
                ("Location:", "")
            ]
            
            for label, value in next_meeting_items:
                p = doc.add_paragraph()
                run1 = p.add_run(f"{label} ")
                run1.font.size = Pt(11)
                run1.font.name = 'Calibri'
                run1.bold = True
                run2 = p.add_run(value)
                run2.font.size = Pt(11)
                run2.font.name = 'Calibri'

            # Save to bytes
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output.getvalue()
            
        except Exception as e:
            print(f"Error during Word document generation: {e}")
            import traceback
            print(traceback.format_exc())
            # Try to save what we have so far if doc exists
            try:
                if 'doc' in locals():
                    output = io.BytesIO()
                    doc.save(output)
                    output.seek(0)
                    print("Saved partial document")
                    return output.getvalue()
            except:
                pass
            raise Exception(f"Failed to generate Word document: {str(e)}")

    def process_content_and_generate_minutes(
        self,
        content: str,
        request: MeetingMinutesRequest
    ) -> bytes:
        """Complete pipeline: Process content with Gemini, generate minutes"""
        try:
            if not content or not content.strip():
                raise Exception("Meeting content cannot be empty")
            
            print(f"Processing {len(content)} characters of meeting content...")
            
            # Process with Gemini (with fallback if it fails)
            print("Processing with Gemini...")
            try:
                processed_data = self.process_content_with_gemini(content)
                print(f"Processed data: {processed_data}")
            except Exception as gemini_error:
                print(f"Gemini processing failed: {gemini_error}")
                # Create fallback structure from content
                print("Using fallback: creating basic structure from content...")
                processed_data = {
                    "meeting_title": request.meeting_title or "Meeting",
                    "agenda_items": [
                        {
                            "item_number": 1,
                            "title": "General Discussion",
                            "description": content[:1000] if len(content) > 0 else "No content provided. Please fill in meeting details manually.",
                            "action_items": []
                        }
                    ],
                    "extracted_date": None,
                    "extracted_location": None,
                    "extracted_company": None
                }
                print("Fallback structure created")
            
            # Update request with extracted data if not provided
            if not request.date_time and processed_data.get("extracted_date"):
                request.date_time = processed_data.get("extracted_date")
            if not request.location and processed_data.get("extracted_location"):
                request.location = processed_data.get("extracted_location")
            if not request.company and processed_data.get("extracted_company"):
                request.company = processed_data.get("extracted_company")
            if processed_data.get("meeting_title") and not request.meeting_title:
                request.meeting_title = processed_data.get("meeting_title")
            
            # Generate Word document
            print("Generating Word document...")
            word_data = self.generate_minutes_word(request, processed_data)
            print(f"Generated Word document: {len(word_data)} bytes")
            
            return word_data
            
        except Exception as e:
            import traceback
            error_trace = traceback.format_exc()
            print(f"Error in process_content_and_generate_minutes: {str(e)}")
            print(f"Traceback: {error_trace}")
            raise Exception(f"Error processing PowerPoint and generating minutes: {str(e)}")
